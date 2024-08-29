import streamlit as st
import requests
import json
import time
from docx import Document
from io import BytesIO

# Configuración de la API de Tune
API_URL = "https://proxy.tune.app/chat/completions"
headers = {
    "Authorization": f"Bearer {st.secrets['TUNE_API_KEY']}",
    "Content-Type": "application/json"
}

def generate_novel_element(prompt, max_tokens=3898):
    data = {
        "model": "meta/llama-3.1-8b-instruct",
        "messages": [
            {"role": "system", "content": "Eres un novelista experto"},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.8,
        "frequency_penalty": 1.04,
        "stream": False
    }

    try:
        response = requests.post(API_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']
    except requests.exceptions.RequestException as e:
        st.error(f"Error en la API: {str(e)}")
        return None

def generate_novel_outline(title, genre, num_chapters):
    prompt = f"Genera una sinopsis, trama, personajes, ambiente, técnica narrativa y tabla de contenidos para una novela titulada '{title}' del género '{genre}'. La tabla de contenidos debe incluir {num_chapters} capítulos con títulos sugerentes."
    return generate_novel_element(prompt)

def generate_chapter(chapter_number, title, genre, synopsis, characters, chapter_title):
    prompt = (f"Escribe el capítulo {chapter_number} titulado '{chapter_title}' de una novela con las siguientes características:\n"
              f"Título: {title}\n"
              f"Género: {genre}\n"
              f"Sinopsis: {synopsis}\n"
              f"Personajes principales: {', '.join(characters)}\n\n"
              f"El capítulo debe ser extenso y detallado. Usa la raya (—) para los diálogos de los personajes.\n\n"
              f"Capítulo {chapter_number}: {chapter_title}\n")
    return generate_novel_element(prompt)

def export_to_docx(title, content):
    doc = Document()
    doc.add_heading(title, 0)

    for line in content.split('\n'):
        if line.startswith('Capítulo'):
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def main():
    st.title("Generador de Novelas")

    if 'novel_outline' not in st.session_state:
        st.session_state['novel_outline'] = None
    if 'novel_chapters' not in st.session_state:
        st.session_state['novel_chapters'] = []

    title = st.text_input("Título de la novela:")
    genre = st.text_input("Género de la novela:")
    num_chapters = st.number_input("Número de capítulos (máximo 24):", min_value=1, max_value=24, value=1)

    if st.button("Generar estructura de la novela"):
        if title and genre:
            with st.spinner("Generando estructura de la novela..."):
                novel_outline = generate_novel_outline(title, genre, num_chapters)
                if novel_outline:
                    st.subheader("Estructura de la novela")
                    st.write(novel_outline)
                    st.session_state['novel_outline'] = novel_outline
        else:
            st.warning("Por favor, introduce el título y el género de la novela.")

    if st.session_state['novel_outline']:
        if st.button("Generar capítulos"):
            progress_bar = st.progress(0)

            chapter_titles = [line.split(': ', 1)[1] for line in st.session_state['novel_outline'].split('\n') if line.startswith('Capítulo')]
            st.session_state['novel_chapters'] = []
            synopsis = st.session_state['novel_outline'].split('Sinopsis:')[1].split('Trama:')[0].strip()
            characters = st.session_state['novel_outline'].split('Personajes:')[1].split('Ambiente:')[0].strip().split(', ')

            for i, chapter_title in enumerate(chapter_titles, 1):
                with st.spinner(f"Generando capítulo {i}: {chapter_title}..."):
                    chapter_content = generate_chapter(i, title, genre, synopsis, characters, chapter_title)
                    if chapter_content:
                        st.session_state['novel_chapters'].append((chapter_title, chapter_content))
                        progress_bar.progress(i / len(chapter_titles))
                    else:
                        st.error(f"No se pudo generar el capítulo {i}: {chapter_title}")
                    time.sleep(1)

            st.success("¡Todos los capítulos han sido generados!")

    if st.session_state['novel_chapters']:
        st.subheader("Capítulos de la novela")
        full_novel_content = st.session_state['novel_outline'] + "\n\n"
        for i, (chapter_title, chapter_content) in enumerate(st.session_state['novel_chapters'], 1):
            with st.expander(f"Capítulo {i}: {chapter_title}"):
                st.write(chapter_content)
            full_novel_content += f"\n\nCapítulo {i}: {chapter_title}\n\n{chapter_content}"

        if st.button("Exportar a DOCX"):
            docx_file = export_to_docx(title, full_novel_content)
            st.download_button(
                label="Descargar DOCX",
                data=docx_file,
                file_name=f"{title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
